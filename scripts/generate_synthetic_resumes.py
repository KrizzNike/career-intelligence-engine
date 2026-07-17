"""
CLI: generate synthetic resumes as PDF + DOCX + a structured JSON sidecar.

Produces balanced, labeled, PII-free resumes for the Week-4 parser and the
Week-6/7 matching/scoring engines.

Usage:
    python scripts/generate_synthetic_resumes.py --per-role 40
    python scripts/generate_synthetic_resumes.py --per-role 100 --seed 7 --out data/raw/resumes/synthetic

Each resume produces THREE files sharing a stem:
    <id>.pdf           human-readable resume (parser input)
    <id>.docx          same content, DOCX variant (parser input)
    <id>.json          ground-truth labels (target_role + canonical skill ids)

A data manifest is (re)written at data/raw/_manifest.csv recording source,
license, count, and generated_at for every batch we produce.
"""
import argparse
import csv
import json
import sys
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from src.config import path  # noqa: E402
from src.data_ingestion.synthetic_resume import (  # noqa: E402
    ResumeGenerator, SyntheticResume,
)

TAXONOMY = path("data", "taxonomy", "skill_taxonomy.yaml")
DEFAULT_OUT = path("data", "raw", "resumes", "synthetic")
MANIFEST = path("data", "raw", "_manifest.csv")

SOURCE = "synthetic_generator"
LICENSE = "CC0-1.0 (synthetic; no real PII)"


# ---------- writers ----------

def write_pdf(resume: SyntheticResume, dest: Path) -> None:
    """Render the resume to a clean one-page PDF via reportlab."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer)

    doc = SimpleDocTemplate(str(dest), pagesize=LETTER,
                            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
                            topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"],
                          fontSize=10, leading=14, spaceAfter=4)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"],
                        fontSize=12, spaceBefore=8, spaceAfter=4)
    name_style = ParagraphStyle("name", parent=styles["Title"],
                                fontSize=16, spaceAfter=2)

    flow = [
        Paragraph(resume.name, name_style),
        Paragraph(f"{resume.email} | {resume.phone} | {resume.location}", body),
        Spacer(1, 6),
        Paragraph("SUMMARY", h2),
        Paragraph(resume.summary, body),
        Paragraph("SKILLS", h2),
        Paragraph(", ".join(resume.skill_names), body),
    ]
    if resume.experience:
        flow.append(Paragraph("EXPERIENCE", h2))
        for job in resume.experience:
            flow.append(Paragraph(
                f"<b>{job['title']}</b> &mdash; {job['company']} "
                f"({job['duration']})", body))
            for b in job["bullets"]:
                flow.append(Paragraph(f"&bull; {b}", body))
    if resume.projects:
        flow.append(Paragraph("PROJECTS", h2))
        for p in resume.projects:
            flow.append(Paragraph(f"<b>{p['name']}</b>", body))
            flow.append(Paragraph(p["description"], body))
    if resume.education:
        flow.append(Paragraph("EDUCATION", h2))
        for e in resume.education:
            flow.append(Paragraph(
                f"{e['degree']} &mdash; {e['school']} ({e['year']})", body))
    if resume.certifications:
        flow.append(Paragraph("CERTIFICATIONS", h2))
        for c in resume.certifications:
            flow.append(Paragraph(f"&bull; {c}", body))
    doc.build(flow)


def write_docx(resume: SyntheticResume, dest: Path) -> None:
    """Render the resume to a .docx via python-docx."""
    import docx
    from docx.shared import Pt
    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(resume.name, level=0)
    doc.add_paragraph(f"{resume.email} | {resume.phone} | {resume.location}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(resume.summary)

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(resume.skill_names))

    if resume.experience:
        doc.add_heading("Experience", level=1)
        for job in resume.experience:
            doc.add_heading(f"{job['title']} — {job['company']}", level=2)
            doc.add_paragraph(job["duration"])
            for b in job["bullets"]:
                doc.add_paragraph(b, style="List Bullet")

    if resume.projects:
        doc.add_heading("Projects", level=1)
        for p in resume.projects:
            doc.add_heading(p["name"], level=2)
            doc.add_paragraph(p["description"])

    if resume.education:
        doc.add_heading("Education", level=1)
        for e in resume.education:
            doc.add_paragraph(f"{e['degree']} — {e['school']} ({e['year']})")

    if resume.certifications:
        doc.add_heading("Certifications", level=1)
        for c in resume.certifications:
            doc.add_paragraph(c, style="List Bullet")

    doc.save(str(dest))


def write_json(resume: SyntheticResume, dest: Path) -> None:
    """Ground-truth labels. This is the evaluation target for Weeks 6/7."""
    payload = {
        "resume_id": resume.resume_id,
        "target_role_id": resume.target_role_id,
        "target_role_name": resume.target_role_name,
        "skills": resume.skills,           # canonical taxonomy ids
        "skill_names": resume.skill_names,
        "name": resume.name,
        "generated_at": resume.generated_at,
        "source": SOURCE,
        "license": LICENSE,
    }
    dest.write_text(json.dumps(payload, indent=2, ensure_ascii=False),
                    encoding="utf-8")


# ---------- manifest ----------

def upsert_manifest(batch_name: str, count: int, out_root: Path) -> None:
    """Insert or update one row in data/raw/_manifest.csv."""
    # out_root may be relative (e.g. from a CLI --out flag); resolve it
    # against CWD before computing the project-relative path for the manifest.
    out_root = (Path.cwd() / out_root).resolve() if not out_root.is_absolute() \
        else out_root.resolve()
    rows: list[dict] = []
    fieldnames = ["batch", "source", "license", "count",
                  "generated_at", "path"]
    if MANIFEST.is_file():
        with MANIFEST.open("r", encoding="utf-8", newline="") as fh:
            rows = list(csv.DictReader(fh))
    rows = [r for r in rows if r["batch"] != batch_name]  # replace if exists
    rows.append({
        "batch": batch_name,
        "source": SOURCE,
        "license": LICENSE,
        "count": str(count),
        "generated_at": datetime.utcnow().isoformat(timespec="seconds") + "Z",
        "path": str(out_root.relative_to(ROOT)),
    })
    MANIFEST.parent.mkdir(parents=True, exist_ok=True)
    with MANIFEST.open("w", encoding="utf-8", newline="") as fh:
        w = csv.DictWriter(fh, fieldnames=fieldnames)
        w.writeheader()
        w.writerows(rows)


# ---------- main ----------

def main() -> int:
    ap = argparse.ArgumentParser(description="Generate synthetic resumes")
    ap.add_argument("--per-role", type=int, default=40,
                    help="resumes per deep role (default 40 -> 200 total)")
    ap.add_argument("--seed", type=int, default=42)
    ap.add_argument("--out", type=Path, default=DEFAULT_OUT)
    ap.add_argument("--formats", nargs="+", default=["pdf", "docx", "json"],
                    choices=["pdf", "docx", "json"])
    args = ap.parse_args()

    args.out.mkdir(parents=True, exist_ok=True)

    gen = ResumeGenerator(taxonomy_path=TAXONOMY, seed=args.seed)
    roles = gen.role_ids
    print(f"Generating {args.per_role} resumes x {len(roles)} roles "
          f"= {args.per_role * len(roles)} total (seed={args.seed})")
    print(f"Roles: {', '.join(roles)}")
    print(f"Output: {args.out}")

    resumes = gen.generate_batch(n_per_role=args.per_role)
    n_files = 0
    for r in resumes:
        stem = args.out / r.resume_id
        if "pdf" in args.formats:
            write_pdf(r, stem.with_suffix(".pdf"))
            n_files += 1
        if "docx" in args.formats:
            write_docx(r, stem.with_suffix(".docx"))
            n_files += 1
        if "json" in args.formats:
            write_json(r, stem.with_suffix(".json"))
            n_files += 1

    total = args.per_role * len(roles)
    upsert_manifest(batch_name="synthetic_resumes_v1",
                    count=total, out_root=args.out)

    print(f"\nDone. {total} resumes, {n_files} files written.")
    print(f"Manifest updated: {MANIFEST.relative_to(ROOT)}")
    # Per-role breakdown
    from collections import Counter
    counts = Counter(r.target_role_id for r in resumes)
    for rid in roles:
        print(f"  {rid:20s} {counts[rid]} resumes")
    return 0


if __name__ == "__main__":
    sys.exit(main())
