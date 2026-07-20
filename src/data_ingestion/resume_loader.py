"""
Resume file loader (Week 4).

Purpose
-------
Read a resume PDF or DOCX file and return a structured intermediate that
the rest of the parsing pipeline consumes, WITHOUT doing any extraction
yet. The intermediate preserves enough structure for the segmenter and
extractors to work for BOTH formats:

    ResumeDocument(
        source_path, file_format, raw_text,         # for PDF heuristic parsing
        blocks: list[Block(text, style, kind)]      # for DOCX style-based parsing
    )

Why a unified intermediate (not "PDF parses one way, DOCX another"):
  The downstream extractors (name, education, experience, skills) should
  operate on a single shape regardless of source. We pay that cost once
  HERE (normalizing blocks/text), so extraction logic stays single-path
  and is unit-testable against either format.

KEY INDUSTRY INSIGHT: DOCX carries semantic styles (Title / Heading 1 /
List Bullet), making segmentation free. PDF loses them, so for PDF we
reconstruct structure from heading patterns (UPPERCASE labels) and
layout heuristics. Both paths feed the SAME block list.

Inputs
------
A path to a .pdf or .docx file (any Path | str).

Outputs
-------
ResumeDocument(dataclass) with raw_text + blocks.

Dependencies
------------
PyMuPDF (fitz), python-docx — both in requirements.txt.

Testing example
---------------
    from src.data_ingestion.resume_loader import load_resume
    doc = load_resume("data/raw/resumes/synthetic/bi_analyst_0000.pdf")
    assert doc.file_format == "pdf"
    assert "Allison Hill" in doc.raw_text

Expected output:
    ResumeDocument(path=..., fmt='pdf', n_blocks=8, raw_text_len=...)
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

# Optional imports resolved at call time so the module imports cleanly even
# if a given format's library isn't installed (helps unit-test isolation).


@dataclass
class Block:
    """A logical paragraph-ish chunk of the resume.

    `style` is the source-native style name (DOCX: 'Title'/'Heading 1'/...;
    PDF: '' since PyMuPDF has no styles). `kind` is a normalized label the
    segmenter uses: 'title' / 'heading' / 'bullet' / 'normal'.
    """
    text: str
    style: str = ""
    kind: str = "normal"


@dataclass
class ResumeDocument:
    source_path: Path
    file_format: str  # "pdf" | "docx"
    raw_text: str
    blocks: list[Block] = field(default_factory=list)

    @property
    def n_blocks(self) -> int:
        return len(self.blocks)


# Section-header labels we recognize after normalization. Synthesized CVs
# use Title Case ("Skills"), publicly scraped PDFs often use UPPERCASE.
# Matching is case-insensitive on a known set so neither style is missed.
SECTION_LABELS = (
    "summary", "objective", "profile", "about",
    "skills", "technical skills", "core competencies",
    "experience", "work experience", "professional experience",
    "employment", "work history",
    "projects", "personal projects", "key projects",
    "education", "academic", "education & training",
    "certifications", "certificates", "licenses",
    "awards", "achievements", "publications",
)


def _classify(text: str, style: str = "") -> str:
    """Map a (text, style) pair to a normalized block kind."""
    if style:
        s = style.lower()
        if "title" in s:
            return "title"
        if "heading 1" in s or s == "heading1":
            return "heading"
        if "heading" in s:  # heading 2/3 -> still a sub-heading
            return "subheading"
        if "list" in s or s in ("list bullet", "list paragraph"):
            return "bullet"
    # PDF has no styles; infer from intrinsic signals.
    stripped = text.strip()
    if not stripped:
        return "normal"
    # Project / job title cues (rendered as Heading 2 in synth, but raw in PDF)
    if stripped.lower().startswith("project:"):
        return "subheading"
    return "normal"


def _load_pdf(path: Path) -> ResumeDocument:
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    parts: list[str] = []
    blocks: list[Block] = []
    for page in doc:
        # get_text("blocks") yields (x0,y0,x1,y1,text,block_no,block_type).
        # We sort by reading order (y then x) so two-column resumes don't
        # interleave columns left-to-right. block_type 1 is image -> skip.
        page_blocks = sorted(
            (b for b in page.get_text("blocks") if b[6] == 0),
            key=lambda b: (round(b[1], 1), round(b[0], 1)),
        )
        for b in page_blocks:
            text = b[4].strip()
            if not text:
                continue
            # Block text from PyMuPDF ends with a newline; collapse internal.
            text = " ".join(text.split())
            parts.append(text)
            blocks.append(Block(text=text, style="", kind=_classify(text)))
    doc.close()
    return ResumeDocument(
        source_path=path, file_format="pdf",
        raw_text="\n".join(parts), blocks=blocks,
    )


def _load_docx(path: Path) -> ResumeDocument:
    import docx  # python-docx

    d = docx.Document(str(path))
    blocks: list[Block] = []
    parts: list[str] = []
    for p in d.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        blocks.append(Block(text=text, style=style,
                            kind=_classify(text, style)))
        parts.append(text)
    return ResumeDocument(
        source_path=path, file_format="docx",
        raw_text="\n".join(parts), blocks=blocks,
    )


_LOADERS = {"pdf": _load_pdf, "docx": _load_docx}


def load_resume(path: str | Path) -> ResumeDocument:
    """Load a resume file -> ResumeDocument. Raises ValueError on unknown ext."""
    p = Path(path)
    ext = p.suffix.lower().lstrip(".")
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"unsupported resume format: {ext} ({p})")
    if not p.exists():
        raise FileNotFoundError(p)
    return loader(p)
